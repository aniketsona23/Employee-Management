from tkinter import *
from tkinter import ttk
from tkcalendar import Calendar
from tkinter import messagebox as msg
import docx
import sql_use
import new_emp


def admin_work():
    # TODO -> Initializing Window
    admin_win = Tk()
    admin_win.title("Admin")
    admin_win.geometry("1200x600")
    admin_win.resizable(0, 1)

    # TODO -> Function for search
    def search():
        to_search = search_val.get().lower()
        category = combo.get()
        search_category = {"Name": 0, "ID": 1, "Phone": 2, "Post": 3, "Entry Time": 4, "Leave Time": 5}
        if to_search:
            category_index = search_category[category]
            tree_ids = tree1.get_children()
            tree_data = list(map(lambda x: tree1.item(x)["values"][category_index], tree_ids))
            tree1.selection_remove(tree1.selection())
            for i in range(len(tree_ids)):
                if to_search in tree_data[i].lower():
                    print("yes")
                    tree1.selection_add(tree_ids[i])

    # TODO -> Making Tab 1
    def tab1_work():
        global tree1
        tree1 = tree_view(tab1, column_dict={"col_num": (1, 2, 3, 4, 5, 6,), "width": (200, 140, 240, 200, 100, 100),
                                             "names": ("Name", "Id", "Phone", "Post", "Enter Time", "Leave Time")},
                          height=300)
        data = sql_use.fetch_data("emp_data")
        for record in data:
            tree1.insert("", END, values=record)
        try:
            tree1.selection_set('I001')
        except TclError:
            pass

    # TODO -> Making Tab 2
    def tab2_work():
        global tree2
        tree2 = tree_view(tab2,
                          column_dict={"col_num": (1, 2, 3, 4, 5, 6), "width": (200, 140, 240, 200, 100, 105),
                                       "names": ("Name", "Id", "Base", "Medical", "Advance", "Bonus")},
                          height=500)

        data = sql_use.fetch_data("emp_salary")
        for record in data:
            tree2.insert("", END, values=record)
        try:
            tree2.selection_set('I001')
        except TclError:
            pass
        Button(tab2, text="Payroll", font=("Bahnschrift semiLight SemiCondensed", 17), command=tab3_work).place(x=400,
                                                                                                                y=300)

    # TODO -> Making Tab 3
    def tab3_work():
        global tree2, net_var
        ntb.select(tab3)

        Label(tab3, text="Name", font=("Bahnschrift semiLight SemiCondensed", 17), relief=GROOVE,
              width=17, bd=3).place(x=100, y=10)
        Label(tab3, text="ID", font=("Bahnschrift semiLight SemiCondensed", 17), relief=GROOVE,
              width=17, bd=3).place(x=100, y=60)
        Label(tab3, text="Post", font=("Bahnschrift semiLight SemiCondensed", 17), relief=GROOVE,
              width=17, bd=3).place(x=100, y=110)
        Label(tab3, text="Total Working Days", font=("Bahnschrift semiLight SemiCondensed", 17), relief=GROOVE,
              width=17, bd=3).place(x=100, y=160)
        Label(tab3, text="Absent", font=("Bahnschrift semiLight SemiCondensed", 17), relief=GROOVE, width=17,
              bd=3).place(x=100, y=210)
        Label(tab3, text="Deduction/day", font=("Bahnschrift semiLight SemiCondensed", 17), relief=GROOVE, width=17,
              bd=3).place(x=100, y=260)
        Label(tab3, text="Total Deductions", font=("Bahnschrift semiLight SemiCondensed", 17), relief=GROOVE, width=17,
              bd=3).place(x=100, y=310)
        Label(tab3, text="Base Salary", font=("Bahnschrift semiLight SemiCondensed", 17), relief=GROOVE, width=17,
              bd=3).place(x=100, y=360)
        Label(tab3, text="Medical Pay", font=("Bahnschrift semiLight SemiCondensed", 17), relief=GROOVE, width=17,
              bd=3).place(x=100, y=410)
        Label(tab3, text="Bonus", font=("Bahnschrift semiLight SemiCondensed", 17), relief=GROOVE, width=17,
              bd=3).place(x=100, y=460)
        Label(tab3, text="Advance", font=("Bahnschrift semiLight SemiCondensed", 17), relief=GROOVE, width=17,
              bd=3).place(x=100, y=510)
        Label(tab3, text="Net Salary", font=("Bahnschrift semibold SemiCondensed", 17), relief=GROOVE, width=17,
              bd=3).place(x=720, y=250)

        cal = Calendar(tab3)
        cal.place(x=700, y=30, width=270, height=200)

        name_var = StringVar()
        id_var = StringVar()
        post_var = StringVar()
        deduct_per_day = IntVar()
        total_deduct = IntVar()
        absent = IntVar()
        base_var = IntVar()
        adv_var = IntVar()
        med_var = IntVar()
        bon_var = IntVar()
        net_var = IntVar()

        name_ent = Entry(tab3, textvariable=name_var, font=("Bahnschrift semiLight SemiCondensed", 17), relief=SUNKEN,
                         state="readonly", bd=3)
        id_ent = Entry(tab3, textvariable=id_var, font=("Bahnschrift semiLight SemiCondensed", 17), relief=SUNKEN,
                       state="readonly", bd=3)
        post_ent = Entry(tab3, textvariable=post_var, font=("Bahnschrift semiLight SemiCondensed", 17), relief=SUNKEN,
                         state="readonly", bd=3)
        total_days_ent = Entry(tab3, font=("Bahnschrift semiLight SemiCondensed", 17), relief=SUNKEN, bd=3)
        absent_ent = Entry(tab3, textvariable=absent, font=("Bahnschrift semiLight SemiCondensed", 17),
                           relief=SUNKEN, bd=3)
        deduct_per_day_ent = Entry(tab3, textvariable=deduct_per_day, font=("Bahnschrift semiLight SemiCondensed", 17),
                                   relief=SUNKEN, bd=3)
        total_deduct_ent = Entry(tab3, textvariable=total_deduct, state="readonly",
                                 font=("Bahnschrift semiLight SemiCondensed", 17),
                                 relief=SUNKEN, bd=3)
        base_ent = Entry(tab3, textvariable=base_var, font=("Bahnschrift semiLight SemiCondensed", 17),
                         relief=SUNKEN, bd=3)
        med_ent = Entry(tab3, textvariable=med_var, font=("Bahnschrift semiLight SemiCondensed", 17),
                        relief=SUNKEN, bd=3)
        adv_ent = Entry(tab3, textvariable=adv_var, font=("Bahnschrift semiLight SemiCondensed", 17),
                        relief=SUNKEN, bd=3)
        bon_ent = Entry(tab3, textvariable=bon_var, font=("Bahnschrift semiLight SemiCondensed", 17),
                        relief=SUNKEN, bd=3)
        net_ent = Entry(tab3, textvariable=net_var, state="readonly", font=("Bahnschrift semiLight SemiCondensed", 17),
                        relief=SUNKEN, bd=3)
        name_ent.place(x=400, y=10)
        id_ent.place(x=400, y=60)
        post_ent.place(x=400, y=110)
        total_days_ent.place(x=400, y=160)
        absent_ent.place(x=400, y=210)
        deduct_per_day_ent.place(x=400, y=260)
        total_deduct_ent.place(x=400, y=310)
        base_ent.place(x=400, y=360)
        med_ent.place(x=400, y=410)
        adv_ent.place(x=400, y=460)
        bon_ent.place(x=400, y=510)
        net_ent.place(x=700, y=300)

        Button(tab3, text='Generate Slip', font=("Bahnschrift semiLight SemiCondensed", 17),
               command=pay_slip).place(x=740, y=350)
        select = tree2.item(tree2.selection())['values']

        emp_pers_data = sql_use.fetch_data("emp_data", condition=f"emp_id='{select[1]}'")[0]
        emp_sal_data = sql_use.fetch_data("emp_salary", condition=f"id='{select[1]}'")[0]
        name_var.set(emp_pers_data[0])
        id_var.set(emp_pers_data[1])
        post_var.set(emp_pers_data[2])
        base_var.set(emp_sal_data[2])
        med_var.set(emp_sal_data[3])
        adv_var.set(emp_sal_data[4])
        bon_var.set(emp_sal_data[5])
        net_sal = base_var.get() + med_var.get() - adv_var.get() + bon_var.get() - total_deduct.get()
        net_var.set(net_sal)
        deduct_per_day_ent.bind("<Key>", lambda event: total_deduct.set(absent.get() * deduct_per_day.get()))
        absent_ent.bind("<Key>", lambda event: total_deduct.set(absent.get() * deduct_per_day.get()))

    notestyle = ttk.Style()

    notestyle.configure("TNotebook.Tab", font=("Bahnschrift semiLight SemiCondensed ", 14), borderwidth=4)
    notestyle.configure("TNotebook", font=("Bahnschrift semiLight SemiCondensed ", 14), background="navy")
    notestyle.configure("Treeview.Heading", font=("Bahnschrift semiLight SemiCondensed ", 14), background="navy")
    notestyle.configure("Treeview",
                        font=("Bahnschrift semiLight SemiCondensed ", 14),
                        background="#E1E1E1", foreground="#000000", rowheight=25)
    notestyle.map('Treeview',
                  background=[('selected', 'dodgerblue3')])
    ntb = ttk.Notebook(admin_win, height=600, width=999, )
    ntb.pack(side=RIGHT)

    tab1 = Frame(ntb, background="navy")
    tab2 = Frame(ntb, background="navy")
    tab3 = Frame(ntb, background="navy")
    tabs = {tab1: "Employee", tab2: "Salary", tab3: "Payroll", }

    for tab in tabs:
        ntb.add(tab, text=tabs[tab])

    paned = PanedWindow(admin_win, width=200, height=600, relief=RAISED, bd=3, bg="navy").pack(side=LEFT)

    new_emp_but = Button(paned, text="New Employee", font=('Bahnschrift semiLight SemiCondensed', 17), bg="#ffee58",
                         fg="navy",
                         command=lambda: new_emp.new_employee((tree1, tree2), admin_win), width=17)
    del_emp_but = Button(paned, text="Delete Employee", font=('Bahnschrift semiLight SemiCondensed', 17), bg="#ffee58",
                         fg="navy", command=delete_emp, width=17)
    update_emp_but = Button(paned, text="Update Employee", font=('Bahnschrift semiLight SemiCondensed', 17),
                            bg="#ffee58", fg="navy",
                            command=lambda: new_emp.new_employee((tree1, tree2), admin_win, True), width=17)

    new_emp_but.place(x=0, y=40)
    del_emp_but.place(x=0, y=90)
    update_emp_but.place(x=0, y=140)

    search_val = StringVar()
    Entry(admin_win, textvariable=search_val, font=('Bahnschrift semiLight SemiCondensed', 17), bd=3,
          width=16).place(x=7, y=454)

    Button(admin_win, text="Search Employee", font=('Bahnschrift semiLight SemiCondensed', 17), bg="#ffee58", fg="navy",
           width=17, command=search).place(x=0, y=500)

    combo = ttk.Combobox(paned, values=("Name", "ID", "Phone", "Post", "Entry Time", "Leave Time"), width=15,
                         font=('Bahnschrift semiLight SemiCondensed', 17),
                         state="readonly", justify=CENTER)
    combo.place(x=8, y=400)
    combo.set("Name")
    ntb.select(tab1)

    tab1_work()
    tab2_work()
    tab3_work()

    ntb.select()
    admin_win.mainloop()


def tree_view(master, column_dict, height):
    tree = ttk.Treeview(master, columns=column_dict["col_num"], height=height, show="headings", )
    scroll = Scrollbar(master, command=tree.yview)
    scroll.pack(side=RIGHT, fill=Y)
    tree.config(yscrollcommand=scroll.set)
    for i in range(len(column_dict["width"])):
        tree.column(i + 1, width=column_dict["width"][i])
        tree.heading(i + 1, text=column_dict["names"][i])
    tree.pack(side=TOP)
    return tree


def delete_emp():
    ans = msg.askyesno("Confirmation", "Are you sure you want to delete record?")
    if ans:
        selected = filter(lambda x: x, (tree1.selection(), tree2.selection()))
        for i in selected:
            value = tree1.item(i)["values"][2]
            sql_use.delete_data("emp_salary", f"id='{value}'")
            sql_use.delete_data("emp_data", f"emp_id='{value}'")
            tree1.delete(i)
            tree2.delete(i)


def pay_slip():
    global data2
    iid_emp = tree2.selection()[0]
    data = tree2.item(iid_emp)["values"]
    data2 = tree1.item(iid_emp)["values"]
    print(data)
    doc = docx.Document()
    doc.add_heading('Pay Slip', 0)
    para = doc.add_paragraph()
    para.add_run(f'Employee Name        : {data[0]}').bold = True
    para.add_run(f'\nEmployee ID               : {data[1]}').bold = True
    para.add_run(f'\nPost\t\t           : {data2[3]}').bold = True

    salaries = ['Basic Salary', 'Bonus', 'Advance', 'Medical Pay', 'Net Salary']
    table = doc.add_table(rows=2, cols=5)
    table.style = 'TableGrid'
    print(table.style)

    for j in range(4):
        cell = table.cell(0, j)
        cell2 = table.cell(1, j)
        cell.text = salaries[j]
        cell2.text = str(data[j + 2])

    cell = table.cell(0, 4)
    cell2 = table.cell(1, 4)
    cell.text = "Net Salary"
    cell2.text = str(net_var.get())
    ask = msg.askquestion('Confirmation', f'Are You sure to make a Payslip for {data[0]}?')
    if ask == 'yes':
        doc.save(f'{data[0]}.docx')
        msg.showinfo('Success', f'Successfully created Payslip for {data[0]}!')
    else:
        pass


if __name__ == '__main__':
    admin_work()
